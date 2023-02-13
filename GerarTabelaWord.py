from docx import Document
import docx.enum.text
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

document = Document()

# Define tamanho da folha como carta
section = document.sections[0]
section.page_width = Cm(21.59)
section.page_height = Cm(27.94)

# Define margens da página
section.left_margin = Cm(0.5)
section.right_margin = Cm(0.5)
section.top_margin = Cm(1.2)
section.bottom_margin = Cm(1.2)

# Adiciona tabela com 3 colunas e 10 linhas
table = document.add_table(rows=10, cols=3)
table.autofit = False
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Define largura e altura das células
for cell in table.columns[0].cells:
    cell.width = Cm(6.7)
    cell.height = Cm(2.5)

for cell in table.columns[1].cells:
    cell.width = Cm(6.7)
    cell.height = Cm(2.5)

for cell in table.columns[2].cells:
    cell.width = Cm(6.7)
    cell.height = Cm(2.5)

# Define espaçamento entre colunas
table.columns[0].width = Cm(6.7 + 0.4)
table.columns[1].width = Cm(6.7)
table.columns[2].width = Cm(6.7 + 0.4)

# Adiciona texto "Teste" em cada célula
for row in table.rows:
    for cell in row.cells:
        para = cell.paragraphs[0]
        run = para.add_run()
        run.text = "Teste"
        run.font.name = "Arial"
        run.font.size = Pt(36)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


# Salva arquivo com o nome "testepdf"
document.save("testepdf.docx")
